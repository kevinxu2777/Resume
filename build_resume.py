from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Jiaqing_Xu_Resume.docx"


def set_font(run, name="Aptos", size=10.0, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_para_spacing(paragraph, before=0, after=2, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_bottom_border(paragraph, color="2F5D62", size="8", space="4"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2F5D62")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_section(doc, title):
    p = doc.add_paragraph()
    set_para_spacing(p, before=7, after=3)
    r = p.add_run(title.upper())
    set_font(r, size=10.5, bold=True, color=(47, 93, 98))
    add_bottom_border(p, size="6", space="2")
    return p


def add_role(doc, left, right, subtitle=None):
    p = doc.add_paragraph()
    set_para_spacing(p, before=3, after=0)
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
    r = p.add_run(left)
    set_font(r, size=10.2, bold=True)
    r = p.add_run("\t" + right)
    set_font(r, size=9.6)
    if subtitle:
        sp = doc.add_paragraph()
        set_para_spacing(sp, before=0, after=1)
        sr = sp.add_run(subtitle)
        set_font(sr, size=9.7, bold=True, color=(70, 70, 70))


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    set_para_spacing(p, before=0, after=1)
    r = p.add_run("- " + text)
    set_font(r, size=9.4)
    return p


def add_skills_line(doc, label, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=0, after=1)
    r = p.add_run(label + ": ")
    set_font(r, size=9.4, bold=True)
    r = p.add_run(text)
    set_font(r, size=9.4)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.45)
    sec.bottom_margin = Inches(0.45)
    sec.left_margin = Inches(0.6)
    sec.right_margin = Inches(0.6)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.4)
    styles["List Bullet"].font.name = "Aptos"
    styles["List Bullet"].font.size = Pt(9.4)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(name, after=1)
    r = name.add_run("Jiaqing Xu")
    set_font(r, size=18, bold=True, color=(31, 47, 52))

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(contact, after=5)
    r = contact.add_run("Greater Boston, MA | kevinxu2777@gmail.com | ")
    set_font(r, size=9.5, color=(55, 55, 55))
    add_hyperlink(contact, "github.com/kevinxu2777", "https://github.com/kevinxu2777")

    summary = doc.add_paragraph()
    set_para_spacing(summary, before=0, after=4, line=1.03)
    r = summary.add_run(
        "Computer Science graduate and junior system administrator with hands-on experience supporting enterprise users, "
        "maintaining operational systems, and building Python automation tools with local dashboards, SQLite persistence, "
        "email alerting, and API integrations."
    )
    set_font(r, size=9.6)

    add_section(doc, "Skills")
    add_skills_line(doc, "Programming", "Python, JavaScript, Java, SQL, Go, HTML, CSS")
    add_skills_line(doc, "Systems and IT", "Windows PowerShell, SecureCRT, VMware, Nutanix, network switch configuration, ticket triage")
    add_skills_line(doc, "Automation and Data", "SQLite, REST APIs, HTML dashboards, SMTP email alerts, macOS LaunchAgents, spreadsheets")
    add_skills_line(doc, "Languages", "English (fluent), Chinese (native)")

    add_section(doc, "Experience")
    add_role(
        doc,
        "HiQ Computers - Onsite at Pfizer, Andover, MA",
        "Sep 2025 - Present",
        "Junior System Administrator",
    )
    add_bullet(doc, "Resolve MCS support tickets for Pfizer end users, balancing daily troubleshooting with operational follow-through.")
    add_bullet(doc, "Maintain and support Non-GMP systems and workflows including ETOP, OMEGA, Unanet, and File Transfer.")
    add_bullet(doc, "Troubleshoot CLAN laptops, Nutanix-related requests, user-reported issues, and basic network/switch tasks.")
    add_bullet(doc, "Execute and validate SPEC-related activities according to documented procedures and support expectations.")
    add_bullet(doc, "Build and maintain operational spreadsheets for issue tracking, reporting, asset visibility, and support coordination.")

    add_section(doc, "Selected Projects")
    add_role(doc, "Market Watch Tool", "Python, SQLite, HTML, SMTP")
    add_bullet(
        doc,
        "Built a local market-monitoring system that tracks global news, macro events, equities, commodities, rates, and volatility signals.",
    )
    add_bullet(
        doc,
        "Implemented event scoring, deduplication, SQLite history, email alert batching, and a live HTML dashboard with filtering and feedback controls.",
    )

    add_role(doc, "Award Travel Copilot", "Python, CLI, SQLite, API integration")
    add_bullet(
        doc,
        "Built a travel-award monitoring workflow that uses seats.aero data to detect new award availability and create actionable email alerts.",
    )
    add_bullet(
        doc,
        "Designed local profile logic for points balances, transfer-partner recommendations, credit tracking, and dashboard-based trip review.",
    )

    add_section(doc, "Education")
    add_role(doc, "Boston University, Boston, MA", "May 2025", "Bachelor of Arts in Computer Science")

    core = doc.core_properties
    core.author = "Jiaqing Xu"
    core.last_modified_by = "Jiaqing Xu"
    core.title = "Jiaqing Xu Resume"
    core.subject = "Resume"
    core.keywords = "resume, systems administration, python, automation"
    core.comments = ""

    doc.save(OUT)


if __name__ == "__main__":
    build()
